"""
文件工具模块，负责批次 ID 生成、文件保存与 docx 解析。
"""
from __future__ import annotations

import re
import random
import string
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, Optional, Tuple

from docx import Document
from fastapi import UploadFile

from config.settings import UPLOAD_DIR
from app.util.logger import logger


SUPPORTED_EXTENSIONS = {".docx", ".md", ".markdown", ".txt"}
INVALID_EXTENSION_MESSAGE = "文件格式错误：仅支持 .docx/.md/.markdown/.txt"
CONTENT_TOO_SHORT_MESSAGE = "正文过短，无法判定有效作业"
PARSE_FAILED_MESSAGE = "无法解析该 Word 文件，可能已损坏"
FILENAME_FORMAT_HINT = "文件命名建议包含班级、姓名、学号、作业名称，例如：25计算机科学与技术1班+张三三+202502210111+职业规划书"
FORMAT_INVALID_MESSAGE = "作业格式不符合要求：正文需宋体小四号，行间距1.5倍"


@dataclass
class FileMeta:
    """从文件名中解析出的基础元信息。"""

    original_name: str
    class_name: Optional[str]
    student_name: Optional[str]
    student_id: Optional[str]
    assignment_title: Optional[str]


def generate_batch_id(prefix: str = "batch") -> str:
    """生成批次 ID，格式为 prefix-年月日时分秒-随机串。"""
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    rand = "".join(random.choices(string.ascii_lowercase + string.digits, k=6))
    return f"{prefix}-{timestamp}-{rand}"


def save_upload_files(batch_id: str, files: Iterable[UploadFile]) -> Tuple[Path, list[Path]]:
    """将上传的文件保存到本地批次目录。"""
    batch_dir = UPLOAD_DIR / batch_id
    batch_dir.mkdir(parents=True, exist_ok=True)
    stored_paths: list[Path] = []
    for file in files:
        target = batch_dir / file.filename
        with target.open("wb") as f:
            content = file.file.read()
            f.write(content)
        logger.info("保存上传文件：%s", target)
        stored_paths.append(target)
    return batch_dir, stored_paths


def extract_student_info(filename: str) -> Tuple[Optional[str], Optional[str]]:
    """
    兼容旧逻辑：仅返回学号与姓名。

    优先按“班级+姓名+学号+作业名称”格式解析，无法识别时退回简单分隔规则。
    """
    meta = parse_filename_meta(filename)
    if meta.student_id or meta.student_name:
        return meta.student_id, meta.student_name

    stem = Path(filename).stem
    if "_" in stem:
        student_id, student_name = stem.split("_", 1)
        return student_id.strip(), student_name.strip()
    if "-" in stem:
        student_id, student_name = stem.split("-", 1)
        return student_id.strip(), student_name.strip()
    return None, None


def parse_filename_meta(filename: str) -> FileMeta:
    """
    解析文件名中的班级、姓名、学号与作业名称。

    规范格式示例：
    25计算机科学与技术1班+张三三+202502210111+职业规划书.docx
    """
    stem = Path(filename).stem.strip()

    def normalize_text(text: str) -> str:
        return (
            text.replace("｜", "|")
            .replace("‖", "|")
            .replace("·", ".")
            .replace("—", "-")
            .replace("–", "-")
            .strip()
        )

    stem_norm = normalize_text(stem)

    def clean_token(token: str) -> str:
        t = token.strip()
        # 去掉常见外壳
        t = re.sub(r"^[【\[\(（]+|[】\]\)）]+$", "", t).strip()
        # 去掉前缀噪音
        t = re.sub(r"^(作业提交|作业|Homework|HOMEWORK|作业_提交)[:：_\-]*", "", t).strip()
        return t

    # 1) 键值对格式：班级=... 姓名=... 学号=... 作业=...
    kv_patterns = {
        "class_name": r"(?:班级|Class|CLASS)",
        "student_name": r"(?:姓名|Name|NAME)",
        "student_id": r"(?:学号|ID|Id|id)",
        "assignment_title": r"(?:作业|Work|WORK|Assignment|ASSIGNMENT)",
    }
    kv_found: dict[str, str] = {}
    for field, key_pat in kv_patterns.items():
        m = re.search(rf"{key_pat}\s*[:=－\-]\s*([^_\-+@|.]+)", stem_norm)
        if m:
            kv_found[field] = clean_token(m.group(1))
    if kv_found.get("student_id") and kv_found.get("student_name"):
        return FileMeta(
            original_name=filename,
            class_name=kv_found.get("class_name") or None,
            student_name=kv_found.get("student_name") or None,
            student_id=kv_found.get("student_id") or None,
            assignment_title=kv_found.get("assignment_title") or None,
        )

    # 2) 统一分隔符切分（支持 + _ - | . @ 空格）
    tokens = [clean_token(t) for t in re.split(r"[+_\-@|.\s]+", stem_norm) if clean_token(t)]

    # 3) 补充“【...】...”这种：将括号内也当作 token
    bracket_tokens = [clean_token(t) for t in re.findall(r"[【\[\(（]([^】\]\)）]+)[】\]\)）]", stem_norm)]
    for bt in bracket_tokens:
        if bt and bt not in tokens:
            tokens.insert(0, bt)

    def guess_student_id(text: str) -> Optional[str]:
        # 学号通常为 6-20 位纯数字
        candidates = re.findall(r"\d{6,20}", text)
        if not candidates:
            return None
        # 取最长的那一个
        return sorted(candidates, key=len, reverse=True)[0]

    def is_class_token(t: str) -> bool:
        return ("班" in t) or re.search(r"\bclass\b", t, flags=re.IGNORECASE) is not None

    def is_name_token(t: str) -> bool:
        # 允许中文姓名、拼音/英文名
        if re.fullmatch(r"[\u4e00-\u9fff·•]{2,10}", t):
            return True
        if re.fullmatch(r"[A-Za-z][A-Za-z.\-]{1,30}", t):
            return True
        return False

    student_id = guess_student_id(stem_norm)
    class_name: Optional[str] = None
    student_name: Optional[str] = None
    assignment_title: Optional[str] = None

    # 4) 从 tokens 中推断字段
    for t in tokens:
        if not t:
            continue
        if student_id is None:
            sid = guess_student_id(t)
            if sid:
                student_id = sid
                continue
        if class_name is None and is_class_token(t):
            class_name = t
            continue
        if student_name is None and is_name_token(t) and (student_id is None or student_id not in t):
            # 避免将任意中文短语误判为姓名：若未识别到学号，则仅在已识别到班级时才接受姓名。
            if student_id is None and class_name is None:
                continue
            student_name = t
            continue

    # 5) 作业名称：尽量取最后一个“不像班级/姓名/学号”的 token
    reserved = set(filter(None, [class_name, student_name, student_id]))
    for t in reversed(tokens):
        if t and t not in reserved:
            assignment_title = t
            break

    # 6) 处理“单 token 但其实是无分隔符拼接”的情况
    if len(tokens) == 1 and student_id and ("班" in stem_norm):
        tokens = []

    # 7) 极端无分隔符：25计科1班张三三202502210111专业分析报告
    if not tokens and (student_id or ("班" in stem_norm)):
        if student_id:
            left, right = stem_norm.split(student_id, 1)
            # 从左边找班级
            m_class = re.search(r"(.+?班)", left)
            if m_class:
                class_name = clean_token(m_class.group(1))
                left_rest = left.replace(m_class.group(1), "")
            else:
                left_rest = left
            # 剩余视为姓名
            name_guess = clean_token(left_rest)
            if name_guess and is_name_token(name_guess):
                student_name = name_guess
            # 右边视为作业名称
            title_guess = clean_token(right)
            if title_guess:
                assignment_title = title_guess

    # 若班级被误判为整串（包含学号/作业名），尝试精炼为“...班”
    if class_name and student_id and student_id in class_name:
        m_refine = re.search(r"(.+?班)", class_name)
        if m_refine:
            class_name = clean_token(m_refine.group(1))

    # 8) 特殊括号拼接格式：【班级】姓名【学号】作业名称
    if student_id and class_name and student_name is None:
        m_name = re.search(r"】\s*([^【\[\(（]{1,20})\s*[【\[\(（]", stem_norm)
        if m_name:
            name_guess = clean_token(m_name.group(1))
            if name_guess and is_name_token(name_guess):
                student_name = name_guess
    if student_id:
        m_title = re.search(r"[】\]\)）]\s*([^【]+)$", stem_norm)
        if m_title:
            title_guess = clean_token(m_title.group(1))
            if title_guess and (assignment_title is None or len(title_guess) < len(assignment_title)):
                assignment_title = title_guess

    # 9) 若完全识别不到学号与姓名，则认为无法解析
    if not student_id and not student_name:
        logger.warning("文件名无法识别出学号与姓名：%s；%s", filename, FILENAME_FORMAT_HINT)
        return FileMeta(
            original_name=filename,
            class_name=None,
            student_name=None,
            student_id=None,
            assignment_title=None,
        )

    # 10) 若仅识别到姓名但无法识别学号，默认不采信（避免误判）
    if student_id is None:
        logger.warning("文件名未识别到学号，已忽略姓名信息：%s；%s", filename, FILENAME_FORMAT_HINT)
        student_name = None
        class_name = None
        assignment_title = None

    return FileMeta(
        original_name=filename,
        class_name=class_name,
        student_name=student_name,
        student_id=student_id,
        assignment_title=assignment_title,
    )


def parse_docx_text(file_path: Path, min_length: int = 50) -> str:
    """解析 docx 正文为纯文本，出错或字数不足时抛出异常。"""
    try:
        document = Document(file_path)
    except Exception as exc:  # noqa: BLE001
        logger.error("解析 docx 失败：%s", exc)
        raise ValueError(PARSE_FAILED_MESSAGE) from exc

    paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
    content = "\n".join(paragraphs)
    if not content or len(content) < min_length:
        raise ValueError(CONTENT_TOO_SHORT_MESSAGE)
    return content


def validate_supported_file(file_path: Path) -> None:
    """校验文件扩展名是否在允许范围内。"""
    if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise ValueError(INVALID_EXTENSION_MESSAGE)


def parse_text_file(file_path: Path, min_length: int = 50) -> str:
    """读取纯文本/Markdown 文件内容，出错或字数不足时抛出异常。"""
    try:
        try:
            content = file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            content = file_path.read_text(encoding="gb18030")
    except Exception as exc:  # noqa: BLE001
        logger.error("读取文本文件失败：%s", exc)
        raise ValueError("无法读取该文本文件，可能编码不受支持或文件已损坏") from exc

    content = content.strip()
    if not content or len(content) < min_length:
        raise ValueError(CONTENT_TOO_SHORT_MESSAGE)
    return content


def parse_file_text(file_path: Path, min_length: int = 50) -> str:
    """根据扩展名解析文件为正文纯文本。"""
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return parse_docx_text(file_path, min_length=min_length)
    return parse_text_file(file_path, min_length=min_length)


def validate_docx_format(
    file_path: Path,
    *,
    allowed_font_keywords: Optional[list[str]] = None,
    allowed_font_size_pts: Optional[list[float]] = None,
    font_size_tolerance: float = 0.5,
    target_line_spacing: Optional[float] = 1.5,
    line_spacing_tolerance: Optional[float] = 0.1,
) -> None:
    """
    校验 docx 正文格式。

    要求：
    1. 文档中至少存在一段“正文”满足：字体为宋体（SimSun/宋体）、字号为小四（约 12pt）、行间距为 1.5 倍；
    2. 不要求全文所有段落都一致，但若全文找不到任何符合规范的正文段落，则判为格式异常。

    若不满足上述“至少一段合规正文”要求，则抛出异常，由上层归类为问题文件。
    """
    # 说明：此校验仅用于内部流程控制，不参与任何提示词构造与大模型输入。
    try:
        document = Document(file_path)
    except Exception as exc:  # noqa: BLE001
        logger.error("读取 docx 失败：%s", exc)
        raise ValueError(PARSE_FAILED_MESSAGE) from exc

    # 默认规则：宋体小四（12pt）、行距 1.5 倍；误差用于兼容不同 Word 环境。
    if not allowed_font_keywords:
        allowed_font_keywords = ["宋体", "SimSun"]
    if not allowed_font_size_pts:
        allowed_font_size_pts = [12.0]

    font_errors: list[str] = []
    size_errors: list[str] = []
    spacing_errors: list[str] = []
    has_valid_body = False

    def get_effective_font_name(run) -> Optional[str]:
        if run.font.name:
            return run.font.name
        if run.style and getattr(run.style, "font", None) and run.style.font.name:
            return run.style.font.name
        return None

    def get_effective_font_size(run) -> Optional[float]:
        if run.font.size:
            return float(run.font.size.pt)
        if run.style and getattr(run.style, "font", None) and run.style.font.size:
            return float(run.style.font.size.pt)
        return None

    def get_effective_line_spacing(para) -> Optional[float]:
        pf = para.paragraph_format
        if pf and pf.line_spacing:
            try:
                return float(pf.line_spacing)
            except Exception:  # noqa: BLE001
                return None
        style_pf = getattr(getattr(para, "style", None), "paragraph_format", None)
        if style_pf and style_pf.line_spacing:
            try:
                return float(style_pf.line_spacing)
            except Exception:  # noqa: BLE001
                return None
        return None

    for index, para in enumerate(document.paragraphs, start=1):
        if not para.text or not para.text.strip():
            continue

        spacing = get_effective_line_spacing(para)
        spacing_ok = True
        if target_line_spacing is not None and line_spacing_tolerance is not None:
            spacing_ok = spacing is not None and abs(spacing - float(target_line_spacing)) <= float(line_spacing_tolerance)
        if not spacing_ok:
            spacing_errors.append(f"第{index}段行距={spacing if spacing is not None else '未设置'}")

        paragraph_has_valid_run = False
        for run in para.runs:
            if not run.text or not run.text.strip():
                continue
            font_name = get_effective_font_name(run)
            font_ok = font_name is not None and any(k in font_name for k in allowed_font_keywords)
            if not font_ok:
                font_errors.append(f"第{index}段字体={font_name if font_name else '未设置'}")

            font_size = get_effective_font_size(run)
            size_ok = False
            if font_size is not None:
                size_ok = any(abs(font_size - float(t)) <= font_size_tolerance for t in allowed_font_size_pts)
            if not size_ok:
                size_errors.append(f"第{index}段字号={font_size if font_size is not None else '未设置'}pt")

            if font_ok and size_ok:
                paragraph_has_valid_run = True

        if spacing_ok and paragraph_has_valid_run:
            has_valid_body = True

    if not has_valid_body:
        details: list[str] = []
        if font_errors:
            details.append("字体问题：" + "；".join(font_errors[:5]))
        if size_errors:
            details.append("字号问题：" + "；".join(size_errors[:5]))
        if spacing_errors:
            details.append("行距问题：" + "；".join(spacing_errors[:5]))
        message = FORMAT_INVALID_MESSAGE + "。"
        if details:
            message += " " + "；".join(details)
        raise ValueError(message)
